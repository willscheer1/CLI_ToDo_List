"""
Create a To-Do list that saves as a text file and has the option of being exported as a.docx file.
"""
import re
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ToDoList():
    
    def __init__(self) -> None:
        """
        Initilizes class instance list with previuosly saved items.
        """
        self.file = "todo_list.txt"
        self.items = self.get_items()

    def get_items(self) -> list[str]:
        """
        Exports saved items from instance text file into the items list.
        If text file does not exist, one is created.

        Returns:
            items (list[str]): List of imported items from text file. If no items
                                in text file, list initializes to [""]. 
        """
        items = [""]    # "" at index 0 acts as placeholder for easier indexing
        try:
            # if file exists, pull items from file
            with open(self.file, "r") as file:
                while item := file.readline():
                    items.append(item.strip())
        except:
            # else create file
            with open(self.file, "w") as file:
                pass

        return items

    def display_list(self) -> None:
        """
        Prints items in the item list to the console.
        """
        print("\nTo-Do List:\n")
        # list is empty
        if len(self.items) == 1:
            print("Your list is empty")
        # iterate through and print items in list
        else:
            for index in range(1, len(self.items)):
                print(f"{index}. {self.items[index].capitalize()}")
        print()
    
    def save_list(self) -> None:
        """
        Saves current item list into the text file.
        """
        with open(self.file, "w") as file:
            for idx in range(1, len(self.items)):
                file.write(f"{self.items[idx]}\n")

    def help(self) -> None:
        """
        Prints list commands to the console.
        """
        print("""Available commands:
    -add: Add's item to the end of the To-Do list. Optional index tag (-i) can be used to specify the index where the item will be added. (ex. add -4 'Do the laundry')
    -del: Delete's first item from the To-Do list. Optional index tag (-i) can be used to specify the index removed. (ex. del -2)
    -clear: Delete's all items from the To-Do list.
    -export: Take's a file path and export's To-Do list as a .docx file. (ex. export your/path/here/my_list)
    -help: Prints list of available commands to the console.
    -exit: Exit the program."""
        )

    def add_item(self, item: str, index: int = 0) -> None:
        """
        Add's given item to the item list at specified index if given, then saves
        the list.

        Parameters:
            item (str): Item to be added to the list.
            index (int): Index where item is to be inserted. 
        """
        if index:
            self.items.insert(index, item)
        else:
            self.items.append(item)

        self.save_list()

    def delete_item(self, index: int = 0) -> None:
        """
        Delete's first item or, if given, item at specified index from the list,
        then saves the list.

        Parameters:
            index (int): Index where item is to be deleted.
        """
        if len(self.items) > 1:
            self.items.pop(index)
            self.save_list()
        else:
            print("No items in list.\n")

    def clear_list(self) -> None:
        """
        Removes all items from the list then saves the list.
        """
        self.items = [self.items[0]]
        self.save_list()

    def export(self, path: str) -> None:
        """
        Exports .docx file to the specified path.

        Parameters:
            path (str): Validated path.
        """
        def format_docx() -> object:
            """
            Creates a .docx document with contents from the To-Do List.

            Returns:
                (obj) Document object containing To-Do list.
            """
            creation_date = date.today().strftime("%m/%d/%Y")
            document = Document()
            # heading
            document.add_heading("To-Do List", level=1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_heading(f"Created on {creation_date}\n", level=2).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # add items
            for item in self.items:
                if item == "":
                    continue
                document.add_paragraph(item.capitalize(), style="List Number")

            return document
        
        document = format_docx()
        try:
            document.save(f"{path}.docx")
        except Exception as e:
            print(f"File failed to save. Error: {e}\n")
        else:
            print("The document saved successfully.\n")

    def get_input(self) -> str:
        """
        Gets list command as user input.

        Returns:
            (str): User input.
        """
        print("Type 'help' to see list of available commands.")
        return input("Command: ")

    def validate_command(self, input_str: str) -> dict:
        """
        Validates command given by user. If error in command,
        prints error to the console.

        Parameters:
            input_str (str): Input string given by user.
        Returns:
            (dict): Dictionary containing keys:
                "isValid" (bool): Bool representing if the command is valid.
                "keyword" (str): Represents type of command given. (ex. 'add')
                ...args (str): Other text given such as index tag
                                or item text, depending on type of command.
        """
        def validate_keyword(keyword: str) -> bool:
            """
            Checks if the given keyword is a valid option.

            Parameters:
                keyword (str): User entered keyword.
            Returns:
                (bool): True if keyword is valid, else False.
            """
            # valid keyword
            if keyword in ["add", "del", "clear", "export", "help", "exit"]:
                return True
            
            # invalid keyword
            print(f"'{keyword}' command not recognizable.\n")
            return False
        
        def index_check(idx: str) -> bool:
            """
            Checks command for an index tag.

            Parameters:
                idx (str): String that my contain an index tag.
            Returns:
                (bool): True if index tag is present, else False.
            """
            if re.match(r"-[0-9]+", idx):
                return True
            return False
        
        def validate_index(idx: int) -> bool:
            """
            Validates the given index is in range of the numbered list items.

            Parameters:
                idx (int): User given index.
            Returns:
                (bool): True if index is in range, else False.
            """
            # index must be in numbered list range.
            if idx >= 1 and idx < len(self.items) or idx == 1:  # if list contains only placeholder, except 1 as index.
                return True
            print(f"Given index of {idx} is out of range.\n")
            return False

        def validate_path(path: str) -> bool:
            """
            Validate's path's character length, naming restrictions, and given file name.
            Does not validate if path exists.

            Parameters:
                path (str): User given path.
            Returns:
                (bool): True if path is considered valid, otherwise False.
            """
            # no empty
            if not path:
                print("'export' command requries a file path.\n")
                return False
            # not longer than 150 chars
            if len(path) > 250:
                print("Path strings are restricted to 250 characters.\n")
                return False
            # missing file name
            if path[-1] == "/" or path[-1] == "\\":
                print("Missing file name from path.\n")
                return False
            # validate file name
            if path.find("/") >= 0:
                file_name = path.split("/")[-1]
            elif path.find("\\") >= 0:
                file_name = path.split("\\")[-1]
            # only file name given (save in current directory)
            else:
                file_name = path
            # no invalid characters (\, /, :, *, ?, ", <, >, |)
            if inv_char := re.search(r"[\\/:*?\"<>|]", file_name):
                print(f"Invalid character '{inv_char.group()}' in file name.\n")
                return False
            
            return True

        split_str = input_str.split(" ")
        # check keyword
        keyword = split_str.pop(0)
        index = 0
        item = ""
        # make sure valid keyword was given
        if not validate_keyword(keyword):
            return {"isValid": False}
        
        # 'exit' command
        if keyword == "exit":
            return {
                "isValid": True,
                "keyword": "exit"
                }

        # 'clear' command
        if keyword == "clear":
            return {
                "isValid": True,
                "keyword": "clear"
            }

        # 'help' command
        if keyword == "help":
            return {
                "isValid": True,
                "keyword": "help"
            }
        
        # 'add' command
        if keyword == "add":
            # only keyword given
            if len(split_str) == 0:
                print("'add' command requires a new list item as input.\n")
                return {"isValid": False}
            # check for index
            if index_check(split_str[0]):
                index = int(split_str.pop(0).strip("-"))
                # validate index
                if not validate_index(index):
                    return {"isValid": False}
            # get given item
            if len(split_str) == 0:
                print("'add' command requires a new list item as input.\n")
                return {"isValid": False}
            else:
                item = " ".join(split_str)

            return {
                "isValid": True,
                "keyword": keyword,
                "index": index,
                "item": item
            }

        # 'del' commmand
        if keyword == "del":
            # check for index
            if len(split_str) > 0 and index_check(split_str[0]):
                index = int(split_str.pop(0).strip("-"))
                # validate index
                if not validate_index(index):
                    return {"isValid": False}
            # don't accept if item given
            if len(split_str) > 0:
                print("'del' command does not accept list item as input.\n")
                return {"isValid": False}
            
            return {
                "isValid": True,
                "keyword": keyword,
                "index": index
            }
        
        # 'export' command
        if keyword == "export":
            path = " ".join(split_str).strip()
            if not validate_path(path):
                return {"isValid": False}
        
            return {
                "isValid": True,
                "keyword": keyword,
                "path": path
            }
            
    def execute_command(self, command: dict) -> None:
        """
        Perform's actions based on given command.

        Parameters:
            command (dict): Dictionary containing command keyword and other 
                            information for action based on keyword.
        """
        keyword = command["keyword"]
        # exit
        if keyword == "exit":
            exit()
        # clear
        if keyword == "clear":
            self.clear_list()
        # help
        elif keyword == "help":
            self.help()
        # add
        elif keyword == "add":
            self.add_item(command["item"], command["index"])
        # delete
        elif keyword == "del":
            self.delete_item(command["index"])
        # export
        elif keyword == "export":
            self.export(command["path"])

    def run(self) -> None:
        """
        Runs the To-Do List program.
        """
        while True:
            # show list
            self.display_list()
            # get input
            command = self.get_input()
            # process command
            request = self.validate_command(command)
            if request["isValid"]:
                self.execute_command(request)



